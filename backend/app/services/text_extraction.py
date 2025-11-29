"""Text extraction service for various document formats"""

import io
import logging
import os
import subprocess
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)


def _configure_tesseract():
    """Auto-configure Tesseract tessdata path if not already set"""
    if os.environ.get('TESSDATA_PREFIX'):
        logger.info(f"Tesseract TESSDATA_PREFIX already set: {os.environ['TESSDATA_PREFIX']}")
        return

    # Try to find tessdata directory (order matters - check version 5 first)
    possible_paths = [
        '/usr/share/tesseract-ocr/5/tessdata',
        '/usr/share/tesseract-ocr/4.00/tessdata',
        '/usr/share/tesseract-ocr/4/tessdata',
        '/usr/share/tessdata',
        '/usr/local/share/tessdata',
    ]

    for path in possible_paths:
        if os.path.exists(path) and os.path.isdir(path):
            # Check if eng.traineddata exists in this directory
            eng_file = os.path.join(path, 'eng.traineddata')
            if os.path.exists(eng_file):
                logger.info(f"Found Tesseract data directory: {path}")
                os.environ['TESSDATA_PREFIX'] = path
                return

    # Fallback: try to set it anyway if directory exists
    for path in possible_paths:
        if os.path.exists(path) and os.path.isdir(path):
            logger.warning(f"Setting TESSDATA_PREFIX to {path} (no eng.traineddata found)")
            os.environ['TESSDATA_PREFIX'] = path
            return

    logger.error("Could not find Tesseract data directory!")

    # Try to get from tesseract command
    try:
        result = subprocess.run(['tesseract', '--version'], capture_output=True, text=True)
        logger.info(f"Tesseract version: {result.stderr.split(chr(10))[0] if result.stderr else 'unknown'}")
    except Exception as e:
        logger.warning(f"Failed to get tesseract version: {e}")


# Configure Tesseract on module import
_configure_tesseract()


class TextExtractionService:
    """Service for extracting text from various document formats"""

    @staticmethod
    async def extract_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file

        Args:
            file_content: PDF file content as bytes

        Returns:
            Extracted text content

        Raises:
            Exception: If PDF extraction fails
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)

            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue

            extracted_text = "\n\n".join(text_content)
            logger.info(f"Extracted {len(extracted_text)} characters from PDF with {len(pdf_reader.pages)} pages")

            return extracted_text

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    async def extract_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file

        Args:
            file_content: DOCX file content as bytes

        Returns:
            Extracted text content

        Raises:
            Exception: If DOCX extraction fails
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            extracted_text = "\n\n".join(text_content)
            logger.info(f"Extracted {len(extracted_text)} characters from DOCX")

            return extracted_text

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    async def extract_from_txt(file_content: bytes, encoding: str = "utf-8") -> str:
        """
        Extract text from plain text file

        Args:
            file_content: Text file content as bytes
            encoding: Text encoding (default: utf-8)

        Returns:
            Extracted text content

        Raises:
            Exception: If text extraction fails
        """
        try:
            # Try different encodings if utf-8 fails
            encodings = [encoding, "utf-8", "latin-1", "cp1252"]

            for enc in encodings:
                try:
                    text = file_content.decode(enc)
                    logger.info(f"Extracted {len(text)} characters from TXT using {enc} encoding")
                    return text
                except UnicodeDecodeError:
                    continue

            raise Exception("Failed to decode text with any supported encoding")

        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise Exception(f"Failed to extract text from TXT: {str(e)}")

    @staticmethod
    async def extract_from_image(file_content: bytes, language: str = "eng") -> str:
        """
        Extract text from image using OCR (Tesseract)

        Args:
            file_content: Image file content as bytes
            language: OCR language (default: eng)

        Returns:
            Extracted text content

        Raises:
            Exception: If OCR fails
        """
        try:
            image = Image.open(io.BytesIO(file_content))

            # Convert to RGB if necessary
            if image.mode not in ("RGB", "L"):
                image = image.convert("RGB")

            # Perform OCR
            text = pytesseract.image_to_string(image, lang=language)

            logger.info(f"Extracted {len(text)} characters from image using OCR")
            return text

        except pytesseract.TesseractNotFoundError:
            logger.error("Tesseract is not installed or not in PATH")
            raise Exception("OCR service not available. Please install Tesseract.")
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            raise Exception(f"Failed to extract text from image: {str(e)}")

    @classmethod
    async def extract_text(
        cls,
        file_content: bytes,
        content_type: str,
        language: Optional[str] = None
    ) -> str:
        """
        Extract text from file based on content type

        Args:
            file_content: File content as bytes
            content_type: MIME type of the file
            language: Language for OCR (optional)

        Returns:
            Extracted text content

        Raises:
            Exception: If extraction fails or content type not supported
        """
        content_type = content_type.lower()

        if "pdf" in content_type:
            return await cls.extract_from_pdf(file_content)

        elif "word" in content_type or "docx" in content_type:
            return await cls.extract_from_docx(file_content)

        elif "text" in content_type or "plain" in content_type:
            return await cls.extract_from_txt(file_content)

        elif "image" in content_type or content_type in [
            "image/png", "image/jpeg", "image/jpg", "image/gif", "image/bmp"
        ]:
            ocr_lang = language or "eng"
            return await cls.extract_from_image(file_content, ocr_lang)

        else:
            raise Exception(f"Unsupported content type: {content_type}")
